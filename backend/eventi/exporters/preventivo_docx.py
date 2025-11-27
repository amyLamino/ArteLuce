# (chemin : /backend/eventi/exporters/preventivo_docx.py)
"""
Export Word du preventivo — version fusionnée et corrigée.

- Compatible avec ton ancien code :
  - expose `render_preventivo_docx(...)`
  - expose `export_preventivo_docx(...)`
- Supporte un template docxtpl (recommandé) avec regroupement par CATEGORIA.
- Fallback python-docx si le template est absent.
- Adapté à tes modèles actuels :
  - Evento (data_evento_da, data_evento_a, note, stato, etc.)
  - RigaEvento (qta, prezzo, importo, copertura_giorni)
  - Materiale (categoria, sottocategoria, nome)
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from datetime import datetime, date
from decimal import Decimal
from collections import defaultdict
from typing import Any, Dict, List

from django.http import HttpResponse
from django.conf import settings
from django.utils.timezone import localtime

# ------------------------------------------------------------
# Dépendances docx / docxtpl
# ------------------------------------------------------------
try:
    from docxtpl import DocxTemplate, InlineImage  # type: ignore
    from docx.shared import Mm  # type: ignore
    HAS_DOCXTPL = True
except Exception:  # pragma: no cover
    HAS_DOCXTPL = False

try:
    import docx  # python-docx
    from docx.shared import Pt  # type: ignore
except Exception:  # pragma: no cover
    docx = None  # type: ignore


# ------------------------------------------------------------
# Helpers de formatage
# ------------------------------------------------------------
def _safe(v: Any, default: str = "") -> str:
    if v is None:
        return default
    s = str(v).strip()
    return s if s else default


def _euro(v: Any) -> str:
    try:
        dv = Decimal(str(v))
    except Exception:
        dv = Decimal("0")
    # format italien simple 1.234,56 €
    s = f"{dv:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    return f"{s} €"


# ------------------------------------------------------------
# Contexte / normalisation
# ------------------------------------------------------------
def _doc_header(evento) -> Dict[str, Any]:
    """
    En-tête documentaire : numéro, révision, date document.
    Utilise les champs éventuels de l'Evento (numero_documento, versione, created_at…)
    sans rien casser si absents.
    """
    num = _safe(getattr(evento, "numero_documento", getattr(evento, "numero", "")), "—")
    # dans ton modèle, tu as "versione"
    rev = getattr(evento, "rev", getattr(evento, "versione", 0)) or 0
    data_doc = getattr(evento, "data_documento", getattr(evento, "created_at", date.today()))
    try:
        dd = localtime(data_doc).strftime("%d/%m/%Y")
    except Exception:
        dd = date.today().strftime("%d/%m/%Y")
    return {"num": num, "rev": rev, "data": dd}


def _emittente_ctx(cfg: Dict[str, Any] | None = None) -> Dict[str, str]:
    """
    Coordonnées de l’émetteur (depuis settings.BRAND par défaut, si présent).
    """
    src = cfg if cfg is not None else getattr(settings, "BRAND", {})
    g = (
        (lambda k: (src.get(k) if isinstance(src, dict) else getattr(src, k, None)) or "—")
        if src
        else (lambda k: "—")
    )

    return {
        "ragsoc": g("ragione_sociale"),
        "indirizzo": g("indirizzo"),
        "cap_citta_prov": g("cap_citta_prov")
        if g("cap_citta_prov") != "—"
        else (
            (g("cap") + " " + g("citta") + " " + g("provincia"))
            if isinstance(src, dict)
            else "—"
        ),
        "cf": g("cf"),
        "piva": g("piva"),
        "rea": g("rea"),
        "referent": g("referent"),
        "phone": g("telefono") if g("telefono") != "—" else g("phone"),
        "email": g("email"),
        "web": g("website") if g("website") != "—" else g("web"),
    }


def _evento_ctx(evento) -> Dict[str, str]:
    """
    Contexte "evento" pour le template :
    - spett: nom du client
    - rif: titre / code offerta
    - data_evento: une date ou un intervalle "dal ... al ..."
    - luogo: nom du lieu
    - stato: stato offerta ou stato interno
    """
    cli = getattr(evento, "cliente", None)
    luogo_obj = getattr(evento, "luogo", None)
    luogo = _safe(getattr(luogo_obj, "nome", getattr(evento, "citta", "")), "—")

    # Gestion de l'intervalle de dates
    data_da = getattr(evento, "data_evento_da", None) or getattr(
        evento, "data_evento", None
    )
    data_a = getattr(evento, "data_evento_a", None) or data_da

    if data_da and data_a:
        if data_da == data_a:
            try:
                data_evento_str = data_da.strftime("%d/%m/%Y")
            except Exception:
                data_evento_str = str(data_da)
        else:
            data_evento_str = (
                f"dal {data_da.strftime('%d/%m/%Y')} al {data_a.strftime('%d/%m/%Y')}"
            )
    else:
        data_ev = getattr(evento, "data_evento", None)
        try:
            data_evento_str = (
                localtime(data_ev).strftime("%d/%m/%Y") if data_ev else "—"
            )
        except Exception:
            data_evento_str = data_ev.strftime("%d/%m/%Y") if data_ev else "—"

    rif = _safe(getattr(evento, "titolo", getattr(evento, "codice_offerta", "")), "—")
    spett = _safe(getattr(cli, "nome", getattr(cli, "ragione_sociale", "")), "—")

    # stato offerta si présent, sinon stato interno
    stato = getattr(evento, "offerta_stato", None) or getattr(evento, "stato", "")

    return {
        "spett": spett,
        "rif": rif,
        "data_evento": data_evento_str,
        "luogo": luogo,
        "stato": stato,
    }


def _build_groups_context(evento) -> Dict[str, Any]:
    """
    Construit:
      - groups: [ { title, rows:[{categoria,nome,qta,pu,tot}], note } ]
      - totals:  récap global

    Adapté à ton modèle :
      - RigaEvento.qta / .prezzo / .importo
      - RigaEvento.materiale.categoria (CharField)
    """
    from ..models import RigaEvento  # import tardif pour éviter les cycles

    righe_qs = (
        RigaEvento.objects.filter(evento=evento)
        .select_related("materiale")
        .order_by("materiale__categoria", "materiale__sottocategoria", "materiale__nome")
    )

    groups_map: Dict[str, List[Dict[str, Any]]] = defaultdict(list)

    tot_materiali = Decimal("0")
    tot_servizi = Decimal("0")
    tot_log = Decimal("0")
    tot_sconti = Decimal("0")

    for r in righe_qs:
        m = r.materiale
        cat = (m.categoria or "").strip() if m else ""
        if not cat:
            cat = "Senza categoria"

        pu = r.prezzo or Decimal("0")
        qta = r.qta or 0
        imp = r.importo or (pu * qta)

        nome = _safe(getattr(m, "nome", ""), "")
        # ajouter info de copertura_giorni si >1
        if getattr(r, "copertura_giorni", 1) and r.copertura_giorni > 1:
            nome += f" (per {r.copertura_giorni} gg)"

        groups_map[cat].append(
            {
                "categoria": cat,
                "nome": nome,
                "qta": qta,
                "pu": _euro(pu),
                "tot": _euro(imp),
            }
        )

        low = (cat or "").lower()
        if low.startswith("serv"):
            tot_servizi += imp
        elif low.startswith("log"):
            tot_log += imp
        else:
            tot_materiali += imp

    groups = [{"title": k, "rows": v, "note": ""} for k, v in groups_map.items()]
    totale = tot_materiali + tot_servizi + tot_log + tot_sconti

    return {
        "groups": groups,
        "totals": {
            "materiali": _euro(tot_materiali),
            "servizi": _euro(tot_servizi),
            "logistica": _euro(tot_log),
            "sconti": _euro(tot_sconti),
            "totale": _euro(totale),
        },
    }


# ------------------------------------------------------------
# Rendu avec template (écriture sur disque)
# ------------------------------------------------------------
def render_preventivo_docx(
    evento, template_path: str, output_path: str, emittente_cfg: Dict[str, Any] | None = None
) -> str:
    """
    Rend un .docx sur disque à partir d’un template docxtpl.
    Si le template n’existe pas, soulève FileNotFoundError (utiliser l’HTTP view qui gère le fallback).
    NE CHANGE PAS l’API attendue par ton code.
    """
    tpl_file = Path(template_path)
    if not tpl_file.exists():
        raise FileNotFoundError(str(tpl_file))

    ctx_groups = _build_groups_context(evento)
    ctx: Dict[str, Any] = {
        "header": _doc_header(evento),
        "emittente": _emittente_ctx(emittente_cfg),
        "evento": _evento_ctx(evento),
        "groups": ctx_groups["groups"],
        "totals": ctx_groups["totals"],
        "note": getattr(evento, "note", "") or "",
    }

    tpl = DocxTemplate(str(tpl_file))
    # logo optionnel
    brand = getattr(settings, "BRAND", {})
    logo_path = (
        Path(brand.get("logo_path", "")) if isinstance(brand, dict) else Path("")
    )
    if HAS_DOCXTPL and logo_path.exists():
        ctx["brand_logo"] = InlineImage(tpl, str(logo_path), width=Mm(35))
    else:
        ctx["brand_logo"] = None

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    tpl.render(ctx)
    tpl.save(output_path)
    return output_path


# ------------------------------------------------------------
# Vue HTTP — export direct au navigateur
# ------------------------------------------------------------
def export_preventivo_docx(
    request, evento_id: int, use_brand_template: bool = True
) -> HttpResponse:
    """
    Vue fonctionnelle compatible avec ton ancien code.
    - Si le template brandé existe + docxtpl dispo → rendu Jinja groupé par catégorie.
    - Sinon → fallback python-docx avec rendu groupé similaire.
    """
    from ..models import Evento  # import tardif

    evento = Evento.objects.select_related("cliente", "luogo").get(pk=evento_id)

    tpl_path = Path("backend/eventi/templates/eventi/preventivo_brand_v1.docx")
    brand = getattr(settings, "BRAND", {})

    buf = BytesIO()

    if use_brand_template and HAS_DOCXTPL and tpl_path.exists():
        # rendu via docxtpl
        try:
            tmp_out = Path("backend/tmp") / f"preventivo_{evento_id}.docx"
            render_preventivo_docx(
                evento,
                str(tpl_path),
                str(tmp_out),
                emittente_cfg=brand if isinstance(brand, dict) else None,
            )
            data = tmp_out.read_bytes()
            tmp_out.unlink(missing_ok=True)
            buf.write(data)
        except Exception:  # pragma: no cover
            # en cas d’erreur de template, on bascule sur fallback
            buf = _fallback_python_docx(evento)
    else:
        buf = _fallback_python_docx(evento)

    filename = f'Preventivo_{evento.id}_{datetime.now().strftime("%Y%m%d_%H%M")}.docx'
    resp = HttpResponse(
        buf.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    resp["Content-Disposition"] = f'attachment; filename="{filename}"'
    return resp


# ------------------------------------------------------------
# Fallback python-docx (sans template) — groupé par catégorie
# ------------------------------------------------------------
def _fallback_python_docx(evento) -> BytesIO:
    """
    Version sans template, utilise python-docx.
    Regroupe par categorie (Materiale.categoria).
    """
    if docx is None:  # pragma: no cover
        raise RuntimeError("python-docx non installato")

    ctx_groups = _build_groups_context(evento)
    buf = BytesIO()

    document = docx.Document()

    # En-tête simple
    head = _doc_header(evento)
    title_run = document.add_paragraph().add_run(
        f'PREVENTIVO n.{head["num"]} rev.{head["rev"]} del {head["data"]}'
    )
    title_run.bold = True
    title_run.font.size = Pt(14)

    ev = _evento_ctx(evento)
    document.add_paragraph(f'Spett. {ev["spett"]}')
    document.add_paragraph(
        f'Rif.: {ev["rif"]}   •   Data evento: {ev["data_evento"]}   •   Luogo: {ev["luogo"]}'
    )
    document.add_paragraph("")

    # Groupes / catégories
    for cat in ctx_groups["groups"]:
        document.add_paragraph("")  # espace
        r = document.add_paragraph().add_run(f'Categoria: {cat["title"]}')
        r.bold = True

        table = document.add_table(rows=1, cols=5)
        hdr = table.rows[0].cells
        hdr[0].text = "Categoria"
        hdr[1].text = "Articolo"
        hdr[2].text = "Q.tà"
        hdr[3].text = "Prezzo unitario"
        hdr[4].text = "Importo"

        for row in cat["rows"]:
            c = table.add_row().cells
            c[0].text = row["categoria"]
            c[1].text = row["nome"]
            c[2].text = str(row["qta"])
            c[3].text = row["pu"]
            c[4].text = row["tot"]

        # Ligne "Nota categoria" fusionnée
        note_row = table.add_row().cells
        merged = note_row[0]
        for i in range(1, 5):
            merged = merged.merge(note_row[i])
        merged.text = "Nota categoria: "

    # Totaux globaux
    t = ctx_groups["totals"]
    document.add_paragraph("")
    document.add_paragraph(
        f'Subtotale materiali: {t["materiali"]}   •   Servizi: {t["servizi"]}   •   Logistica: {t["logistica"]}'
    )
    tot_run = document.add_paragraph(f'TOTALE: {t["totale"]}').runs[0]
    tot_run.bold = True

    # Note globale si présente
    note = _safe(getattr(evento, "note", ""))
    if note:
        document.add_paragraph("")
        document.add_paragraph("Note generali:")
        document.add_paragraph(note)

    document.save(buf)
    buf.seek(0)
    return buf
