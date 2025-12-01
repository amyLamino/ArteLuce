# backend/eventi/views.py
from __future__ import annotations

from datetime import date, datetime, timedelta
from decimal import Decimal
from collections import OrderedDict, defaultdict
import io
import os
import re
import logging
from pathlib import Path

from .views_history import create_revision_if_changed

from rest_framework.viewsets import ModelViewSet

from django.conf import settings
from django.db import transaction
from django.db.models import Max, Prefetch, Q
from django.http import HttpResponse, JsonResponse, HttpResponseRedirect
from django.utils.dateparse import parse_date

from rest_framework import status, viewsets, permissions, generics
from rest_framework.decorators import action, api_view, permission_classes
from rest_framework.permissions import IsAuthenticatedOrReadOnly
from rest_framework.response import Response
from rest_framework.views import APIView

from .views_history import create_revision_if_changed
from .serializers import EventoSerializer

from .models import (
    Cliente, Luogo, Materiale,
    Evento, RigaEvento, CalendarioSlot,
    EventoRevision, MaterialeSuggerito, RegolaSuggerimento,
    Tecnico, Mezzo,
)
from .serializers import (
    ClienteSerializer, LuogoSerializer, MaterialeSerializer,
    EventoSerializer, RigaEventoSerializer, EventoRevisionSerializer,
    TecnicoSerializer, MezzoSerializer
)

logger = logging.getLogger(__name__)

# -------------------------------------------------------------------
# Helpers
# -------------------------------------------------------------------

# chemin : backend/eventi/views.py
from django.http import JsonResponse

def home(request):
    """
    Petite home JSON pour l'API.
    Accessible sur /api/
    """
    return JsonResponse(
        {
            "name": "ArteLuce API",
            "version": "1.0",
            "status": "ok",
        }
    )

def _money_s(x) -> str:
    try:
        q = Decimal(str(x or 0)).quantize(Decimal("0.01"))
    except Exception:
        q = Decimal("0.00")
    s = f"{q:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    return f"{s} €"

def _date_s(d) -> str:
    if not d:
        return ""
    if isinstance(d, (date, datetime)):
        return d.strftime("%d/%m/%Y")
    try:
        return date.fromisoformat(str(d)).strftime("%d/%m/%Y")
    except Exception:
        return str(d)

def _month_bounds(yyyy_mm: str):
    y, m = map(int, yyyy_mm.split("-"))
    start = date(y, m, 1)
    nxt = (start.replace(day=28) + timedelta(days=4)).replace(day=1)
    end = nxt - timedelta(days=1)
    return start, end

def daterange(start: date, end: date):
    """Génère toutes les dates entre start et end (inclus)."""
    cur = start
    while cur <= end:
        yield cur
        cur += timedelta(days=1)

# -------------------------------------------------------------------
# Disponibilité du calendrier
# -------------------------------------------------------------------

class CalendarioAvailability(APIView):
    permission_classes = [permissions.AllowAny]

    def get(self, request):
        """GET /api/calendario/availability?data=YYYY-MM-DD"""
        dstr = request.query_params.get("data")
        try:
            d = date.fromisoformat(str(dstr))
        except Exception:
            return Response({"detail": "Param 'data' manquant ou invalide (YYYY-MM-DD)."}, status=400)

        used = set(
            CalendarioSlot.objects
            .filter(data=d)
            .values_list("location_index", flat=True)
        )
        all_slots = set(range(1, 9))  # 1..8
        free = sorted(list(all_slots - used))
        suggested = free[0] if free else None
        return Response({"used": sorted(list(used)), "free": free, "suggested": suggested})

# -------------------------------------------------------------------
# Anagrafica & Matériaux
# -------------------------------------------------------------------

class ClienteViewSet(viewsets.ModelViewSet):
    queryset = Cliente.objects.all().order_by("-id")
    serializer_class = ClienteSerializer
    permission_classes = [permissions.AllowAny]

class LuogoViewSet(viewsets.ModelViewSet):
    queryset = Luogo.objects.all().order_by("-id")
    serializer_class = LuogoSerializer
    permission_classes = [permissions.AllowAny]

class MaterialeViewSet(viewsets.ModelViewSet):
    queryset = Materiale.objects.all().order_by("nome")
    serializer_class = MaterialeSerializer
    permission_classes = [permissions.AllowAny]

class TecnicoViewSet(MaterialeViewSet):
    def get_queryset(self):
        return super().get_queryset().filter(is_tecnico=True)
    def perform_create(self, serializer):
        serializer.save(is_tecnico=True, is_messo=False, scorta=0)

class MezzoViewSet(MaterialeViewSet):
    def get_queryset(self):
        return super().get_queryset().filter(is_messo=True)
    def perform_create(self, serializer):
        serializer.save(is_messo=True, is_tecnico=False, scorta=1)

# -------------------------------------------------------------------
# Evento + Révisions
# -------------------------------------------------------------------

class EventoViewSet(viewsets.ModelViewSet):
    queryset = Evento.objects.all().order_by("-data_evento", "-id")
    serializer_class = EventoSerializer

    """
    Endpoints utilisés par le front :
      - GET   /api/eventi/?month=YYYY-MM
      - GET   /api/eventi/<id>/
      - PATCH/PUT /api/eventi/<id>/
      - GET/POST /api/eventi/<id>/revisions/
      - GET   /api/eventi/<id>/history/
      - GET   /api/eventi/<id>/audit/
      - PUT   /api/eventi/<id>/righe/
      - GET   /api/eventi/<id>/docx/
      - GET   /api/eventi/next-slot?date=YYYY-MM-DD
    """
    serializer_class = EventoSerializer
    permission_classes = [permissions.AllowAny]
    lookup_value_regex = r"\d+"

    def get_queryset(self):
        qs = (
            Evento.objects
            .select_related("cliente", "luogo")
            .prefetch_related(Prefetch("righe", queryset=RigaEvento.objects.select_related("materiale")))
            .order_by("-id")
        )
        month = self.request.query_params.get("month")
        if month and re.match(r"^\d{4}-\d{2}$", month):
            start, end = _month_bounds(month)
            qs = qs.filter(data_evento__range=(start, end))
        return qs

    # -----------------------------------------------------------------
    # Ajout dynamique des champs de stock sur la sortie JSON
    # -----------------------------------------------------------------
    def _attach_stock_fields(self, data, queryset):
        """
        Ajoute pour chaque evento:
          - stock_tot_scorta : somme des scorte des matériels utilisés
          - stock_tot_dispon : scorta_tot - qta_totale (sur cet evento)
        ➜ utilisé par la liste mensile + EventStockBadge.
        """
        ev_ids = [getattr(ev, "id", None) for ev in queryset]
        ev_ids = [i for i in ev_ids if i is not None]
        if not ev_ids:
            for row in data:
                row["stock_tot_scorta"] = 0
                row["stock_tot_dispon"] = 0
            return data

        righe = (
            RigaEvento.objects
            .select_related("materiale")
            .filter(evento_id__in=ev_ids)
        )

        stats = defaultdict(lambda: {"scorta": 0, "booked": 0})
        seen_materials = defaultdict(set)

        for r in righe:
            eid = r.evento_id
            s = stats[eid]
            s["booked"] += int(r.qta or 0)

            mid = r.materiale_id
            if mid not in seen_materials[eid]:
                sc = int(getattr(r.materiale, "scorta", 0) or 0)
                s["scorta"] += sc
                seen_materials[eid].add(mid)

        for row in data:
            eid = row.get("id")
            st = stats.get(eid)
            if not st:
                row["stock_tot_scorta"] = 0
                row["stock_tot_dispon"] = 0
            else:
                sc = st["scorta"]
                disp = max(0, sc - st["booked"])
                row["stock_tot_scorta"] = sc
                row["stock_tot_dispon"] = disp

        return data

    def list(self, request, *args, **kwargs):
        """
        Surcharge de la liste pour ajouter stock_tot_scorta / stock_tot_dispon
        sans modifier le serializer existant.
        """
        queryset = self.filter_queryset(self.get_queryset())

        page = self.paginate_queryset(queryset)
        if page is not None:
            serializer = self.get_serializer(page, many=True)
            data = self._attach_stock_fields(serializer.data, page)
            return self.get_paginated_response(data)

        serializer = self.get_serializer(queryset, many=True)
        data = self._attach_stock_fields(serializer.data, queryset)
        return Response(data)

    # -----------------------------------------------------------------
    # Hooks de création / modification
    # -----------------------------------------------------------------
    def perform_create(self, serializer):
        evento = serializer.save()
        create_revision_if_changed(evento, note="Creazione evento")

    def perform_update(self, serializer):
        evento = serializer.save()
        create_revision_if_changed(evento, note="Modifica evento")

    @action(detail=True, methods=["get", "post"], url_path="revisions")
    def revisions(self, request, pk=None):
        evento = self.get_object()

        if request.method.lower() == "get":
            qs = evento.revisions.order_by("ref", "created_at")
            return Response(EventoRevisionSerializer(qs, many=True).data)

        note = request.data.get("note") or ""

        rev = create_revision_if_changed(evento, note=note)

        if rev is None:
            last = evento.revisions.order_by("-ref").first()
            if last is None:
                return Response(
                    {"detail": "Nessuna modifica da salvare."},
                    status=status.HTTP_200_OK,
                )
            return Response(
                EventoRevisionSerializer(last).data,
                status=status.HTTP_200_OK,
            )

        return Response(
            EventoRevisionSerializer(rev).data,
            status=status.HTTP_201_CREATED,
        )

    @action(detail=True, methods=["get"], url_path="history")
    def history(self, request, pk=None):
        evento = self.get_object()
        qs = evento.revisions.order_by("ref", "created_at")
        return Response(EventoRevisionSerializer(qs, many=True).data)

    @action(detail=True, methods=["get"], url_path="audit")
    def audit(self, request, pk=None):
        evento = self.get_object()
        qs = evento.revisions.order_by("ref", "created_at")
        return Response(EventoRevisionSerializer(qs, many=True).data)

    @action(detail=True, methods=["put"], url_path="righe")
    def replace_righe(self, request, pk=None):
        ev = self.get_object()
        righe = request.data.get("righe", []) or []
        with transaction.atomic():
            ev.righe.all().delete()
            bulk = []
            for r in righe:
                bulk.append(RigaEvento(
                    evento=ev,
                    materiale_id=int(r["materiale"]),
                    qta=int(r.get("qta", 1) or 1),
                    prezzo=Decimal(str(r.get("prezzo", 0) or 0)),
                    importo=Decimal(str(r.get("qta", 1) or 1)) * Decimal(str(r.get("prezzo", 0) or 0)),
                    is_tecnico=bool(r.get("is_tecnico", False)),
                    is_trasporto=bool(r.get("is_trasporto", False)),
                    copertura_giorni=int(r.get("copertura_giorni", 1) or 1),
                ))
            RigaEvento.objects.bulk_create(bulk)

        last = ev.revisions.aggregate(m=Max("ref"))["m"] or 0
        snap = EventoSerializer(ev, context={"request": request}).data
        EventoRevision.objects.create(evento=ev, ref=last + 1, note="Replace righe", payload=snap)

        return Response(EventoSerializer(ev, context={"request": request}).data)

    # --- Export DOCX --------------------------------------------------


    @action(detail=True, methods=["get"], url_path="docx")
    def export_docx(self, request, pk=None):
        # Imports « souples »
        try:
            from docxtpl import DocxTemplate, InlineImage
        except Exception:
            DocxTemplate = None
            InlineImage = None
        try:
            from docx import Document
            from docx.shared import Mm
        except Exception:
            Document = None
            Mm = None

        ev = self.get_object()

        # ---------- lignes + totaux ---------- #
        rows, imponibile = [], Decimal("0.00")
        pos = 1
        for r in ev.righe.select_related("materiale").all():
            q = Decimal(str(getattr(r, "qta", 0) or 0))
            pu = Decimal(str(getattr(r, "prezzo", 0) or 0))
            imp = (q * pu).quantize(Decimal("0.01"))

            mat = getattr(r, "materiale", None)
            nome = getattr(mat, "nome", None) or f"#{getattr(r, 'materiale_id', '')}"
            cat = (getattr(mat, "categoria", "") or "— Senza categoria —").strip()
            sub = (getattr(mat, "sottocategoria", "") or "— Senza sottocategoria —").strip()

            rows.append(
                {
                    "pos": pos,
                    "categoria": cat,
                    "sottocategoria": sub,
                    "materiale_nome": nome,
                    "qta": int(q),
                    "prezzo": pu,
                    "prezzo_s": f"{pu:.2f} €",
                    "importo": imp,
                    "importo_s": f"{imp:.2f} €",
                }
            )
            imponibile += imp
            pos += 1

        iva_pct = Decimal("22")
        iva = (imponibile * iva_pct / Decimal("100")).quantize(Decimal("0.01"))
        totale = (imponibile + iva).quantize(Decimal("0.01"))
        acconto = Decimal(str(getattr(ev, "acconto_importo", 0) or 0)).quantize(
            Decimal("0.01")
        )
        saldo = max(Decimal("0.00"), totale - acconto)

        # ---------- regroupement (cat -> sub -> items) ---------- #
        by_cat: "OrderedDict[str, OrderedDict[str, list]]" = OrderedDict()
        for line in rows:
            by_cat.setdefault(line["categoria"], OrderedDict())
            by_cat[line["categoria"]].setdefault(line["sottocategoria"], [])
            by_cat[line["categoria"]][line["sottocategoria"]].append(line)

        groups = [
            {"cat": c, "subs": [{"sub": s, "items": items} for s, items in subs.items()]}
            for c, subs in by_cat.items()
        ]

        table_rows = []
        for c, subs in by_cat.items():
            table_rows.append({"kind": "cat", "label": c})
            for s, items in subs.items():
                table_rows.append({"kind": "sub", "label": s})
                for it in items:
                    table_rows.append(
                        {
                            "kind": "item",
                            "pos": it["pos"],
                            "materiale_nome": it["materiale_nome"],
                            "qta": it["qta"],
                            "prezzo": it["prezzo_s"],
                            "importo": it["importo_s"],
                        }
                    )

        # ---------- contexte ---------- #
        co = getattr(settings, "COMPANY", {}) or {}
        ctc = co.get("contact", {}) or {}
        note_raw = (getattr(ev, "note", "") or "").strip()
        cat_notes = getattr(ev, "categoria_notes", {}) or {}

        ctx = {
            # en-tête & méta
            "id": ev.id,
            "versione": getattr(ev, "versione", 0),
            "oggi": date.today(),
            "titolo": getattr(ev, "titolo", "") or "",
            "data_evento": getattr(ev, "data_evento", None),

            # destinataires
            "cliente_nome": getattr(getattr(ev, "cliente", None), "nome", "") or "",
            "luogo_nome": getattr(getattr(ev, "luogo", None), "nome", "") or "",

            # totaux stringifiés
            "imponibile_s": f"{imponibile:.2f} €",
            "iva_pct": f"{iva_pct}",
            "iva_s": f"{iva:.2f} €",
            "totale_ivato_s": f"{totale:.2f} €",
            "acconto_importo_s": f"{acconto:.2f} €",
            "saldo_s": f"{saldo:.2f} €",

            # société
            "emit_ragione": co.get("ragione", ""),
            "emit_indirizzo": co.get("indirizzo", ""),
            "emit_cap_citta": co.get("cap_citta", ""),
            "emit_cf": co.get("cf", ""),
            "emit_piva": co.get("piva", ""),
            "emit_cciaa": co.get("cciaa", ""),

            # contact
            "contatto_nome": ctc.get("nome", ""),
            "contatto_ruolo": ctc.get("ruolo", ""),
            "contatto_cell": ctc.get("cell", ""),
            "contatto_email": ctc.get("email", ""),
            "contatto_web": ctc.get("web", ""),

            # données détaillées
            "righe": rows,
            "groups": groups,
            "table_rows": table_rows,

            # notes
            "note_generali": note_raw,
            "categoria_notes_items": sorted(
                [(k, (v or "").strip()) for k, v in cat_notes.items() if (v or "").strip()],
                key=lambda x: x[0].lower(),
            ),
        }

        candidates = [
            Path(settings.BASE_DIR)
            / "backend"
            / "eventi"
            / "templates"
            / "eventi"
            / "preventivo.docx",
            Path(settings.BASE_DIR)
            / "eventi"
            / "templates"
            / "eventi"
            / "preventivo.docx",
            Path(__file__).resolve().parent / "templates" / "eventi" / "preventivo.docx",
        ]
        tpl_path = next((str(p) for p in candidates if os.path.exists(p)), None)

        # ---------- docxtpl avec filtres date/datefmt ---------- #
        if DocxTemplate and tpl_path:
            tpl = DocxTemplate(tpl_path)

            # ➜ essayer de récupérer un env Jinja existant, sinon en créer un
            try:
                # certaines versions de docxtpl
                env = tpl.get_jinja_env()  # peut lever AttributeError
            except Exception:
                env = getattr(tpl, "environment", None) or getattr(
                    tpl, "jinja_env", None
                )
            if env is None:
                # fallback : nouvel Environment Jinja "simple"
                from jinja2 import Environment

                env = Environment(autoescape=True)

            # --- filtres date / datefmt utilisés dans le modèle ---
            def _date_filter(value, fmt="%d/%m/%Y"):
                if not value:
                    return ""
                if isinstance(value, (datetime, date)):
                    d = value
                else:
                    try:
                        d = date.fromisoformat(str(value))
                    except Exception:
                        return str(value)
                return d.strftime(fmt)

            env.filters["date"] = _date_filter
            env.filters["datefmt"] = _date_filter

            # --- filtre monnaie : {{ x|money }} -> "123,45 €" ---
            def _money_filter(value):
                if value is None or value == "":
                    return ""
                try:
                    q = Decimal(str(value or 0)).quantize(Decimal("0.01"))
                except Exception:
                    return str(value)
                # format "123,45 €" façon italienne
                s = f"{q:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
                return f"{s} €"

            env.filters["money"] = _money_filter
            # alias éventuel si tu utilises aussi |euro dans le modèle
            env.filters["euro"] = _money_filter

            # logo (facultatif)
            logo_path = co.get("logo_path")
            if logo_path and os.path.exists(logo_path) and InlineImage:
                try:
                    ctx["logo"] = InlineImage(tpl, logo_path, width=Mm(28)) if Mm else None
                except Exception:
                    pass

            # bloc "evento" pour le template
            ctx["evento"] = {
                "id": ev.id,
                "versione": getattr(ev, "versione", 0) or 0,
                "titolo": getattr(ev, "titolo", "") or "",
                "data_evento": getattr(ev, "data_evento", None),
                "cliente": {
                    "nome": getattr(getattr(ev, "cliente", None), "nome", "") or ""
                },
                "luogo": {"nome": getattr(getattr(ev, "luogo", None), "nome", "") or ""},
                "righe": [
                    {
                        "pos": it["pos"],
                        "categoria": it["categoria"],
                        "sottocategoria": it["sottocategoria"],
                        "articolo": it["materiale_nome"],
                        "materiale_nome": it["materiale_nome"],
                        "qta": it["qta"],
                        "prezzo": it["prezzo_s"],
                        "importo": it["importo_s"],
                    }
                    for it in rows
                ],
                "totali": {
                    "imponibile": ctx["imponibile_s"],
                    "iva_pct": ctx["iva_pct"],
                    "iva": ctx["iva_s"],
                    "totale": ctx["totale_ivato_s"],
                    "acconto": ctx["acconto_importo_s"],
                    "saldo": ctx["saldo_s"],
                },
                "table_rows": ctx.get("table_rows", []),
            }

            # ➜ très important : on passe *notre* env à render
            tpl.render(ctx, jinja_env=env)

            buf = io.BytesIO()
            tpl.save(buf)
            buf.seek(0)

            # plus besoin d'import ici, date et datetime sont déjà importés en haut

            date_part = ""
            d = ctx.get("data_evento")
            if d:
                if isinstance(d, (datetime, date)):
                    date_part = d.strftime("%Y%m%d")
                else:
                    try:
                        date_part = date.fromisoformat(str(d)).strftime("%Y%m%d")
                    except Exception:
                        date_part = str(d).replace("/", "-").replace(" ", "_")

            filename = f"preventivo_{ev.id}"
            if date_part:
                filename += f"_{date_part}"
            filename += ".docx"

            return HttpResponse(
                buf.getvalue(),
                content_type=(
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                ),
                headers={
                    "Content-Disposition": f'attachment; filename="preventivo_{ev.id}.docx"'
                },
            )

        # ---------- fallback python-docx ---------- #
        if Document:
            doc = Document()
            doc.add_heading(f"Preventivo evento #{ev.id}", 0)
            if ctx["titolo"]:
                doc.add_paragraph(f"Titolo: {ctx['titolo']}")
            if ctx["data_evento"]:
                doc.add_paragraph(
                    "Data: "
                    + _date_filter(ctx["data_evento"])  # même formatage que le filtre
                )
            if ctx["cliente_nome"]:
                doc.add_paragraph(f"Cliente: {ctx['cliente_nome']}")
            if ctx["luogo_nome"]:
                doc.add_paragraph(f"Luogo: {ctx['luogo_nome']}")

            t = doc.add_table(rows=1, cols=6)
            h = t.rows[0].cells
            h[0].text, h[1].text, h[2].text, h[3].text, h[4].text, h[5].text = (
                "POS",
                "CATEGORIA",
                "ARTICOLO",
                "Qtà",
                "PU",
                "Importo",
            )
            for it in rows:
                c = t.add_row().cells
                c[0].text = str(it["pos"])
                c[1].text = it["categoria"]
                c[2].text = it["materiale_nome"]
                c[3].text = str(it["qta"])
                c[4].text = it["prezzo_s"]
                c[5].text = it["importo_s"]

            for line in (
                    f"Imponibile: {ctx['imponibile_s']}",
                    f"IVA {ctx['iva_pct']}%: {ctx['iva_s']}",
                    f"Totale: {ctx['totale_ivato_s']}",
                    f"Acconto: {ctx['acconto_importo_s']}",
                    f"Saldo: {ctx['saldo_s']}",
            ):
                doc.add_paragraph(line)

            if note_raw:
                doc.add_paragraph("")
                doc.add_paragraph("Note generali:")
                for ln in note_raw.splitlines():
                    if ln.strip():
                        doc.add_paragraph(ln, style=None)

            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return HttpResponse(
                buf.getvalue(),
                content_type=(
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                ),
                headers={
                    "Content-Disposition": f'attachment; filename="preventivo_{ev.id}.docx"'
                },
            )

        return Response({"error": "DOCX non supportato."}, status=501)

    # --- Slot disponible pour une date --------------------------------

    @action(detail=False, methods=["get"], url_path="next-slot")
    def next_slot(self, request):
        date_str = request.query_params.get("date", "")
        d = parse_date(date_str)
        if not d:
            return Response({"slot": 1})
        taken = set(
            Evento.objects.filter(data_evento=d).values_list("location_index", flat=True)
        )
        for s in range(1, 9):
            if s not in taken:
                return Response({"slot": s})
        return Response({"slot": 1})

# -------------------------------------------------------------------
# Listes mensuelles (calendario/lista)
# -------------------------------------------------------------------
from django.db.models import Q

class EventiMensiliView(APIView):
    """GET /api/eventi/mese?year=YYYY&month=M -> liste compacte des événements du mois"""
    permission_classes = [permissions.AllowAny]

    def get(self, request):
        y = int(request.query_params.get("year", date.today().year))
        m = int(request.query_params.get("month", date.today().month))

        # bornes du mois
        month_start = date(y, m, 1)
        next_month = (month_start.replace(day=28) + timedelta(days=4)).replace(day=1)
        month_end = next_month - timedelta(days=1)

        qs = (
            Evento.objects
            .filter(
                # 1) anciens événements simples sur data_evento
                Q(data_evento__range=(month_start, month_end))
                |
                # 2) événements multi-jours qui se chevauchent avec le mois
                Q(data_evento_da__lte=month_end, data_evento_a__gte=month_start)
            )
            .select_related("cliente")
            .order_by("data_evento", "location_index", "id")
        )

        data = []
        for ev in qs:
            # pour le calendrier on garde un champ "data_evento" = début
            start = getattr(ev, "data_evento_da", None) or ev.data_evento
            data.append({
                "id": ev.id,
                "titolo": ev.titolo,
                "data_evento": start.isoformat() if start else None,
                "data_evento_da": getattr(ev, "data_evento_da", None) and ev.data_evento_da.isoformat(),
                "data_evento_a": getattr(ev, "data_evento_a", None) and ev.data_evento_a.isoformat(),
                "location_index": ev.location_index,
                "stato": ev.stato,
                "cliente_nome": ev.cliente.nome if ev.cliente_id else None,
                "offerta_stato": getattr(ev, "offerta_stato", None),
                "saldo_state": getattr(ev, "saldo_state", None),
            })

        return Response(data, status=200)


@api_view(["GET"])
@permission_classes([IsAuthenticatedOrReadOnly])
def eventi_mensili(request):
    month = request.query_params.get("month")  # "YYYY-MM"
    if not month or not re.match(r"^\d{4}-\d{2}$", month):
        return Response([])
    start, end = _month_bounds(month)
    qs = (
        Evento.objects
        .filter(data_evento__range=(start, end))
        .select_related("cliente")
        .order_by("data_evento", "location_index")
    )
    data = [{
        "id": ev.id,
        "titolo": ev.titolo,
        "data_evento": ev.data_evento.isoformat(),
        "location_index": ev.location_index,
        "stato": ev.stato,
        "cliente_nome": ev.cliente.nome if ev.cliente_id else None,
    } for ev in qs]
    return Response(data)

def eventi_mensili_alias(request):
    month = request.GET.get("month")
    if not month:
        return JsonResponse({"detail": "Param 'month' required (YYYY-MM)."}, status=400)
    try:
        y, m = month.split("-", 1)
        m = str(int(m))
    except Exception:
        return JsonResponse({"detail": "Bad 'month' format. Expected YYYY-MM."}, status=400)
    return HttpResponseRedirect(f"/api/eventi/mese?year={y}&month={m}")

# -------------------------------------------------------------------
# Pricing & Suggestions (offerta rapida)
# -------------------------------------------------------------------

class PricingView(APIView):
    permission_classes = [permissions.AllowAny]

    def post(self, request):
        """
        JSON: {"materiale_id": int, "luogo_id": int, "qta": int, "data": "YYYY-MM-DD"}
        """
        try:
            mat = Materiale.objects.get(pk=int(request.data["materiale_id"]))
            luogo = Luogo.objects.get(pk=int(request.data["luogo_id"]))
            qta = int(request.data["qta"])
            _ = date.fromisoformat(str(request.data["data"]))
        except Exception as e:
            return Response({"detail": f"Parametri invalidi: {e}"}, status=400)

        totale = (Decimal(mat.prezzo_base or 0) * Decimal(qta))
        return Response({"totale": f"{Decimal(totale):.2f}"}, status=200)

class SuggestionView(APIView):
    permission_classes = [permissions.AllowAny]

    def get(self, request):
        """
        GET /api/suggest?materiale=<id>&date=YYYY-MM-DD

        Logique :
        - suggestions "liées" via MaterialeSuggerito (comme avant)
        - + AUTOSUGGEST :
            • autres produits de la même sottocategoria
            • si rien → autres produits de la même categoria
        """
        try:
            mat_id = int(request.query_params.get("materiale"))
            day = request.query_params.get("date")
            trigger = Materiale.objects.get(pk=mat_id)
        except Exception:
            return Response({"items": []})

        out: list[dict] = []
        seen: set[int] = set()

        # petit helper de disponibilité (tu pourras l'améliorer plus tard
        # avec l'API /magazzino/bookings si tu veux)
        def availability(m: Materiale):
            pren = 0  # pour l'instant on ne recalcule pas ici
            disp = max(0, (m.scorta or 0) - pren)
            return int(pren), disp

        # ---------------------------------------------------------
        # a) Liens directs existants (MaterialeSuggerito)
        # ---------------------------------------------------------
        for link in (
            MaterialeSuggerito.objects
            .filter(trigger=trigger, active=True)
            .select_related("suggested")
        ):
            m = link.suggested
            if m.id in seen or m.id == trigger.id:
                continue
            seen.add(m.id)
            pren, disp = availability(m)
            out.append({
                "materiale_id": m.id,
                "nome": m.nome,
                "prezzo": float(m.prezzo_base or 0),
                "qty_default": link.qty_default or 1,
                "label": link.label or "",
                "is_tecnico": m.is_tecnico,
                "is_messo": m.is_messo,
                "prenotato": pren,
                "disponibilita": disp,
            })

        # ---------------------------------------------------------
        # b) AUTOSUGGEST : même sotto-categoria
        # ---------------------------------------------------------
        same_sub_qs = (
            Materiale.objects
            .filter(
                is_archived=False,
                categoria=trigger.categoria,
                sottocategoria=trigger.sottocategoria,
            )
            .exclude(pk=trigger.pk)
            .order_by("nome")
        )[:10]

        for m in same_sub_qs:
            if m.id in seen:
                continue
            seen.add(m.id)
            pren, disp = availability(m)
            out.append({
                "materiale_id": m.id,
                "nome": m.nome,
                "prezzo": float(m.prezzo_base or 0),
                "qty_default": 1,
                "label": "Stessa sottocategoria",
                "is_tecnico": m.is_tecnico,
                "is_messo": m.is_messo,
                "prenotato": pren,
                "disponibilita": disp,
            })

        # ---------------------------------------------------------
        # c) AUTOSUGGEST : même categoria (si sottocategoria vide
        #    ou si on veut proposer encore d'autres choses)
        # ---------------------------------------------------------
        if not same_sub_qs.exists():
            same_cat_qs = (
                Materiale.objects
                .filter(
                    is_archived=False,
                    categoria=trigger.categoria,
                )
                .exclude(pk=trigger.pk)
                .order_by("nome")
            )[:10]

            for m in same_cat_qs:
                if m.id in seen:
                    continue
                seen.add(m.id)
                pren, disp = availability(m)
                out.append({
                    "materiale_id": m.id,
                    "nome": m.nome,
                    "prezzo": float(m.prezzo_base or 0),
                    "qty_default": 1,
                    "label": "Stessa categoria",
                    "is_tecnico": m.is_tecnico,
                    "is_messo": m.is_messo,
                    "prenotato": pren,
                    "disponibilita": disp,
                })

        # ---------------------------------------------------------
        # d) Bonus tecnici / mezzi (comme avant)
        # ---------------------------------------------------------
        for m in (
            Materiale.objects
            .filter(is_tecnico=True, is_archived=False)
            .order_by("nome")[:2]
        ):
            if m.id in seen:
                continue
            seen.add(m.id)
            pren, disp = availability(m)
            out.append({
                "materiale_id": m.id,
                "nome": m.nome,
                "prezzo": float(m.prezzo_base or 0),
                "qty_default": 1,
                "label": "Tecnico",
                "is_tecnico": True,
                "is_messo": False,
                "prenotato": pren,
                "disponibilita": disp,
            })

        for m in (
            Materiale.objects
            .filter(is_messo=True, is_archived=False)
            .order_by("nome")[:2]
        ):
            if m.id in seen:
                continue
            seen.add(m.id)
            pren, disp = availability(m)
            out.append({
                "materiale_id": m.id,
                "nome": m.nome,
                "prezzo": float(m.prezzo_base or 0),
                "qty_default": 1,
                "label": "Trasporto / Messo",
                "is_tecnico": False,
                "is_messo": True,
                "prenotato": pren,
                "disponibilita": disp,
            })

        return Response({"items": out})

# -------------------------------------------------------------------
# Magazzino (état & réservations)
# -------------------------------------------------------------------
from rest_framework.decorators import api_view, permission_classes
from rest_framework import permissions
# (ces imports tu les as déjà, mais vérifie qu’ils sont bien présents)

# -------------------------------------------------------------------
# Magazzino – calendrier annuel (pour Sinottico)
# -------------------------------------------------------------------

@api_view(["GET"])
@permission_classes([permissions.AllowAny])
def magazzino_calendar(request):
    """
    GET /api/magazzino/calendar?year=2025

    Réponse:
    {
      "year": 2025,
      "days": ["2025-01-01", ...],
      "materials": [
        { "id": 1, "nome": "...", "categoria": "...", "scorta": 10 },
        ...
      ],
      "bookings": [
        { "materiale": 1, "date": "2025-01-05", "qta": 3 },
        ...
      ]
    }

    ➜ ici on tient compte de TOUTE la durée de l'événement
       (data_evento_da / data_evento_a ou data_evento simple).
    """
    # --- année demandée ---
    try:
        year = int(request.GET.get("year") or date.today().year)
    except ValueError:
        year = date.today().year

    start = date(year, 1, 1)
    end   = date(year, 12, 31)

    # --- liste de tous les jours de l'année ---
    days = []
    cur = start
    while cur <= end:
        days.append(cur.isoformat())
        cur += timedelta(days=1)

    # --- matériaux non archivés ---
    mats_qs = (
        Materiale.objects
        .filter(is_archived=False)
        .order_by("categoria", "nome")
    )
    materials = [
        {
            "id": m.id,
            "nome": m.nome,
            "categoria": m.categoria or "— Senza categoria —",
            "scorta": int(m.scorta or 0),
        }
        for m in mats_qs
    ]

    # --- carte (date_iso, materiale_id) -> qta totale ce jour-là ---
    per_day: dict[tuple[str, int], int] = {}

    righe = (
        RigaEvento.objects
        .select_related("evento")
        .filter(
            Q(evento__data_evento_da__lte=end, evento__data_evento_a__gte=start)
            | Q(evento__data_evento__range=(start, end))
        )
        .exclude(evento__stato="annullato")
    )

    for r in righe:
        ev = r.evento

        # support data_evento_da / data_evento_a OU data_evento
        ev_start = getattr(ev, "data_evento_da", None) or ev.data_evento
        ev_end   = getattr(ev, "data_evento_a", None) or ev.data_evento
        if not ev_start or not ev_end:
            continue

        # intersection avec l'année demandée
        s = max(ev_start, start)
        e = min(ev_end, end)
        if e < s:
            continue

        q = int(r.qta or 0)
        cur = s
        while cur <= e:
            key = (cur.isoformat(), r.materiale_id)
            per_day[key] = per_day.get(key, 0) + q
            cur += timedelta(days=1)

    bookings = [
        {"materiale": mid, "date": d, "qta": q}
        for (d, mid), q in per_day.items()
    ]

    return Response(
        {
            "year": year,
            "days": days,
            "materials": materials,
            "bookings": bookings,
        }
    )

from django.db.models import Q
from datetime import timedelta
# et tu as déjà daterange défini plus haut dans le fichier

class MagazzinoStatusView(APIView):
    """
    GET /api/magazzino/status?from=YYYY-MM-DD&to=YYYY-MM-DD&materials=1,2,3

    ➜ tient compte de toute la durée de l'événement
       (data_evento_da / data_evento_a ou data_evento + copertura_giorni)
    """
    permission_classes = [permissions.AllowAny]

    def get(self, request):
        dfrom_str = request.GET.get("from") or request.GET.get("date") or request.GET.get("day")
        dto_str   = request.GET.get("to")   or dfrom_str
        if not dfrom_str:
            return Response({"detail": "Missing 'from' or 'date' parameter."}, status=400)
        try:
            dfrom = date.fromisoformat(dfrom_str)
            dto   = date.fromisoformat(dto_str)
        except ValueError:
            return Response({"detail": "Bad date format. Use YYYY-MM-DD."}, status=400)
        if dto < dfrom:
            dto = dfrom

        ids_param = request.GET.get("materials") or request.GET.get("material")
        if ids_param:
            mids = [int(x) for x in str(ids_param).split(",") if x.strip()]
            mats = Materiale.objects.filter(id__in=mids)
        else:
            mats = Materiale.objects.all()

        # liste des jours de l'intervalle
        days: list[date] = []
        cur = dfrom
        while cur <= dto:
            days.append(cur)
            cur += timedelta(days=1)

        # Evénements qui SE CHEVAUCHENT avec [dfrom, dto]
        righe = (
            RigaEvento.objects
            .select_related("materiale", "evento")
            .filter(
                Q(evento__data_evento_da__lte=dto, evento__data_evento_a__gte=dfrom)
                | Q(evento__data_evento__gte=dfrom, evento__data_evento__lte=dto)
            )
            .exclude(evento__stato="annullato")
        )
        if ids_param:
            righe = righe.filter(materiale_id__in=[m.id for m in mats])

        # carte matériau → date → qta
        by_mat_day: dict[int, dict[date, int]] = {
            m.id: {d: 0 for d in days} for m in mats
        }

        for r in righe:
            ev = r.evento
            # support data_evento_da / data_evento_a OU data_evento + copertura_giorni
            ev_start = getattr(ev, "data_evento_da", None) or ev.data_evento
            ev_end   = getattr(ev, "data_evento_a", None)
            if not ev_end:
                giorni = max(1, int(getattr(r, "copertura_giorni", 1) or 1))
                ev_end = ev_start + timedelta(days=giorni - 1)

            # intersection avec l'intervalle demandé
            s = max(ev_start, dfrom)
            e = min(ev_end, dto)
            if e < s:
                continue

            q = int(r.qta or 0)
            cur = s
            while cur <= e:
                if r.materiale_id in by_mat_day:
                    by_mat_day[r.materiale_id][cur] += q
                cur += timedelta(days=1)

        out = {
            "days": [d.isoformat() for d in days],
            "materials": [],
        }
        for m in mats:
            row = {
                "id": m.id,
                "nome": m.nome,
                "stock": int(m.scorta or 0),
                "by_day": [],
            }
            for d in days:
                used = by_mat_day[m.id][d]
                free = max(0, row["stock"] - used)
                if row["stock"] == 0:
                    status_s = "ko"
                elif free == 0 and row["stock"] > 0:
                    status_s = "warn"
                else:
                    status_s = "ok"
                row["by_day"].append(
                    {"date": d.isoformat(), "used": used, "free": free, "status": status_s}
                )
            out["materials"].append(row)

        return Response(out)



class MagazzinoBookingsView(APIView):
    """
    GET /api/magazzino/bookings?material=<id>&from=YYYY-MM-DD&to=YYYY-MM-DD&on=YYYY-MM-DD

    - calcule les quantités réservées pour CHAQUE jour de l'intervalle [from, to]
    - renvoie :
        scorta        = stock du matériel
        prenotato     = qté réservée le jour "on" (ou le max sur l'intervalle)
        prenotato_max = max des réservations sur la période
        disponibile   = min(scorta - prenotato_giorno) sur toutes les dates
        per_day       = détail par jour
        rows          = liste des eventi impliqués
    """
    permission_classes = [permissions.AllowAny]

    def get(self, request):
        material_id = request.query_params.get("material")
        if not material_id:
            return Response({"error": "query param 'material' mancante"}, status=400)

        try:
            materiale = Materiale.objects.get(pk=int(material_id))
        except Materiale.DoesNotExist:
            return Response({"error": "Materiale non trovato"}, status=404)

        from_s = request.query_params.get("from")
        to_s = request.query_params.get("to")
        on_s = request.query_params.get("on")

        today = date.today()
        d_from = parse_date(from_s) or today
        d_to = parse_date(to_s) or d_from
        if d_to < d_from:
            d_to = d_from

        d_on = parse_date(on_s) if on_s else None

        # Evénements qui se chevauchent avec l’intervalle demandé
        righe = (
            RigaEvento.objects.filter(materiale_id=material_id)
            .select_related("evento", "evento__cliente")
            .filter(
                Q(evento__data_evento_da__lte=d_to, evento__data_evento_a__gte=d_from)
                | Q(evento__data_evento__gte=d_from, evento__data_evento__lte=d_to)
            )
            .filter(evento__stato__in=["bozza", "confermato", "fatturato"])
        )

        per_day: dict[date, int] = {}
        rows = []

        for r in righe:
            ev: Evento = r.evento
            ev_start = getattr(ev, "data_evento_da", None) or ev.data_evento
            ev_end = getattr(ev, "data_evento_a", None) or ev.data_evento
            if not ev_start or not ev_end:
                continue

            # intersection avec l’intervalle demandé
            start = max(ev_start, d_from)
            end = min(ev_end, d_to)
            if end < start:
                continue

            # ligne descriptive par evento
            rows.append({
                "evento_id": ev.id,
                "titolo": ev.titolo,
                "stato": ev.stato,
                "cliente": getattr(ev.cliente, "nome", None),
                "data_evento_da": ev_start.isoformat(),
                "data_evento_a": ev_end.isoformat(),
                "qta": int(r.qta or 0),
                "location_index": ev.location_index,
            })

            # on compte la qté sur chaque jour couvert par l’événement
            for g in daterange(start, end):
                per_day[g] = per_day.get(g, 0) + int(r.qta or 0)

        scorta = int(materiale.scorta or 0)
        if per_day:
            prenotato_max = max(per_day.values())
            disp_min = min(scorta - q for q in per_day.values())
        else:
            prenotato_max = 0
            disp_min = scorta

        # pour compatibilité avec l’UI : "Prenotato (ON)"
        if d_on is not None:
            pren_on = per_day.get(d_on, 0)
        else:
            pren_on = prenotato_max

        return Response(
            {
                "materiale": materiale.id,
                "scorta": scorta,
                "prenotato": pren_on,
                "prenotato_max": prenotato_max,
                "disponibile": max(0, disp_min),
                "per_day": {
                    d.isoformat(): per_day.get(d, 0)
                    for d in daterange(d_from, d_to)
                },
                "rows": rows,
            }
        )
